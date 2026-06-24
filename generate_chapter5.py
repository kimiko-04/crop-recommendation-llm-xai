"""
Generates Chapter 5: Conclusion and Recommendations as a .docx file.
Structure follows UiTM thesis skeleton: Revisit Objectives, Strengths,
Limitations, Further Works, Summary — aligned with 3 objectives from Chapter 1.
Run: python generate_chapter5.py
Output: Chapter5_Conclusion_and_Future_Work.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Style helpers ─────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=12, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading1(text):
    """e.g. 5.1 REVISIT OBJECTIVES"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, bold=True, size=12)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p


def heading2(text):
    """e.g. 5.1.1 FIRST OBJECTIVE"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, bold=True, size=12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.5)
    return p


# ════════════════════════════════════════════════════════════════════════════════
# CHAPTER TITLE
# ════════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CHAPTER 5")
set_font(run, bold=True, size=16)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CONCLUSION AND RECOMMENDATIONS")
set_font(run, bold=True, size=16)
p.paragraph_format.space_after = Pt(20)

# ── Chapter intro ─────────────────────────────────────────────────────────────
body(
    "This chapter evaluates how well every objective of this project was achieved by "
    "revisiting each objective. It also explores and highlights the strengths as well as the "
    "limitations of the developed crop recommendation system and makes recommendations for "
    "further works that can be done to improve the system, before summarising the key points "
    "discussed at the end."
)

# ════════════════════════════════════════════════════════════════════════════════
# 5.1 REVISIT OBJECTIVES
# ════════════════════════════════════════════════════════════════════════════════

heading1("5.1 REVISIT OBJECTIVES")

body(
    "This section will revisit each of the objectives of the project to evaluate how well "
    "each objective was achieved. The project aims to develop a crop recommendation system "
    "that integrates Large Language Model (LLM) and Explainable Artificial Intelligence (XAI) "
    "techniques to produce accurate crop predictions alongside human-readable explanations. "
    "Thus, the project objectives include:"
)

# Objectives list — indented, no bullet, matching the template style
for letter, text in [
    ("a.", "To identify and assess suitable Large Language Model (LLM) and Explainable AI "
           "(XAI) techniques for crop recommendation systems."),
    ("b.", "To design and develop the crop recommendation system using LLM and XAI."),
    ("c.", "To evaluate the functionality and performance of the suggested Crop "
           "Recommendation System."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.75)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(letter + " ")
    set_font(r1, size=12)
    r2 = p.add_run(text)
    set_font(r2, size=12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

doc.add_paragraph()

# ── 5.1.1 ─────────────────────────────────────────────────────────────────────

heading2("5.1.1 FIRST OBJECTIVE")

body(
    "The first objective of this project was to identify and assess suitable Large Language "
    "Model (LLM) and Explainable AI (XAI) techniques for the crop recommendation system. "
    "This objective was addressed during the early phases of the project, particularly in "
    "Chapter 2, which is the Literature Review chapter. Various machine learning and deep "
    "learning models were evaluated, with a focus on models that can process structured "
    "agricultural data effectively while producing interpretable outputs. Traditional models "
    "such as Random Forest, Support Vector Machine, and Decision Tree were considered alongside "
    "transformer-based models. After the evaluation process, BERT (Bidirectional Encoder "
    "Representations from Transformers) and its lighter variant DistilBERT were selected as "
    "the LLM component of the system because of their established performance in sequence "
    "classification tasks and their ability to capture contextual relationships within the "
    "structured natural language input derived from the field parameters."
)

body(
    "For the XAI component, two complementary methods were assessed. SHAP (SHapley Additive "
    "exPlanations) was selected as the model-agnostic feature attribution method because it "
    "provides theoretically grounded numerical explanations for individual predictions. The "
    "BERT CLS-token attention mechanism was selected as a second XAI technique because it "
    "offers a lightweight, natural language reasoning output that is accessible to non-technical "
    "users such as farmers and agricultural officers. This objective is therefore successfully "
    "achieved, as the most suitable LLM and XAI techniques were identified through the "
    "literature review and comparative assessment conducted in Chapter 2."
)

# ── 5.1.2 ─────────────────────────────────────────────────────────────────────

heading2("5.1.2 SECOND OBJECTIVE")

body(
    "The second objective is to design and develop the crop recommendation system using the "
    "identified LLM and XAI techniques. This objective is successfully achieved through the "
    "full development of the CropAI system, as discussed in Chapter 3, which is the Research "
    "Methodology chapter. The system introduces a tabular-to-text conversion pipeline that "
    "transforms the seven numerical field parameters — Nitrogen, Phosphorus, Potassium, pH, "
    "Temperature, Humidity, and Rainfall — into structured natural language sentences. These "
    "sentences are then tokenised and passed to the fine-tuned BERT and DistilBERT sequence "
    "classifiers, which output a probability distribution across 22 crop classes."
)

body(
    "The XAI layer was developed on top of the classifier pipeline. The SHAP KernelExplainer "
    "was integrated to compute per-feature attribution values, and the attention reasoning "
    "module was implemented to extract the CLS-token attention weights and generate a "
    "plain-English explanation sentence identifying the most influential factors behind the "
    "recommendation. The full system was deployed as a web application using FastAPI for the "
    "backend, React with Tailwind CSS for the frontend, and MongoDB for data persistence. "
    "The web application provides user authentication, a crop prediction dashboard, prediction "
    "history, and a comprehensive admin portal for model management, user management, and "
    "data drift monitoring. This objective is therefore fully achieved, as the complete system "
    "was designed, developed, and deployed as described in Chapter 3."
)

# ── 5.1.3 ─────────────────────────────────────────────────────────────────────

heading2("5.1.3 THIRD OBJECTIVE")

body(
    "The third objective is to evaluate the functionality and performance of the suggested "
    "Crop Recommendation System. This objective is achieved in Chapter 4, which is the "
    "Results and Discussion chapter, through both quantitative performance evaluation and "
    "functional testing of all system features. The model performance was evaluated on a "
    "held-out test set derived from the Kaggle crop recommendation dataset, which contains "
    "2,200 samples across 22 crop classes. BERT achieved a classification accuracy of 98.18% "
    "and DistilBERT achieved 97.50%, demonstrating that both transformer models are highly "
    "effective classifiers for the text-serialised crop recommendation task."
)

body(
    "In addition to the quantitative benchmarking, the SHAP feature attribution values and "
    "the attention-based reasoning outputs were validated by comparing the feature rankings "
    "produced by each method against established agronomic knowledge. The directional "
    "agreement observed between SHAP rankings and attention rankings across multiple sample "
    "predictions provides confidence that the model has learned meaningful agricultural "
    "feature associations. The functional testing carried out on the user interface confirmed "
    "that all system features, including registration, login, prediction, SHAP explanation, "
    "attention reasoning, prediction history, and admin management, operate correctly and as "
    "intended. This objective is therefore successfully achieved, as both the performance and "
    "the functionality of the system were thoroughly evaluated in Chapter 4."
)

# ════════════════════════════════════════════════════════════════════════════════
# 5.2 STRENGTHS
# ════════════════════════════════════════════════════════════════════════════════

heading1("5.2 STRENGTHS")

body(
    "One strength of the CropAI system is that it provides dual-layer explainability alongside "
    "its crop predictions. Unlike conventional crop recommendation systems that only output a "
    "predicted crop class, this system produces both a SHAP feature attribution chart and a "
    "natural language reasoning sentence for each prediction. This allows users to understand "
    "not only what the system recommends but also why, which increases transparency and user "
    "trust. The combination of a numerical explanation for technically-oriented users and a "
    "plain-English explanation for non-technical users makes the system accessible to a broader "
    "range of stakeholders, from agricultural researchers to smallholder farmers."
)

body(
    "Another strength of the system is its comprehensive admin portal, which supports the full "
    "model lifecycle without requiring any code changes. Administrators can retrain a new model "
    "candidate in the background, track retraining progress in real time, evaluate the "
    "candidate against the current production model, promote it with temperature calibration, "
    "and monitor data drift using PSI-based statistical analysis across all seven input "
    "features. This design ensures that the system can be maintained and improved over time "
    "as new agricultural data becomes available, making it suitable for long-term deployment "
    "beyond the academic prototype stage."
)

# ════════════════════════════════════════════════════════════════════════════════
# 5.3 LIMITATIONS
# ════════════════════════════════════════════════════════════════════════════════

heading1("5.3 LIMITATIONS")

body(
    "The first limitation of the CropAI system is the high computational cost of the SHAP "
    "explanation. The SHAP KernelExplainer is model-agnostic and treats the BERT model as a "
    "black box, which requires approximately 700 forward passes per explanation request. This "
    "results in a computation time of approximately 30 seconds on CPU hardware, which may be "
    "impractical for users who require immediate feedback or for deployment on low-resource "
    "devices. Various optimisation strategies such as reducing the number of SHAP background "
    "samples were considered, but a more fundamental solution would require replacing the "
    "KernelExplainer with a gradient-based attribution method that is natively compatible "
    "with the PyTorch model."
)

body(
    "The second limitation is the fixed and rule-based nature of the tabular-to-text "
    "conversion template used to serialise field parameters into natural language sentences. "
    "The qualitative descriptors generated by this template, such as 'high' or 'low', are "
    "determined by comparing input values against training-set means and do not account for "
    "regional or seasonal agricultural norms. This limits the representativeness of the "
    "generated text for inputs that fall significantly outside the training distribution "
    "and may affect the quality of the attention-based reasoning output in edge cases."
)

body(
    "The third limitation is that all seven field parameters must be entered manually by the "
    "user through the web interface. The system does not currently support live data "
    "acquisition from IoT soil sensors or external weather APIs, which means the system "
    "depends entirely on the user providing accurate and up-to-date values for all parameters. "
    "This manual input requirement may reduce the practicality of the system in real-world "
    "field conditions where automated data collection would be preferred."
)

# ════════════════════════════════════════════════════════════════════════════════
# 5.4 FURTHER WORKS
# ════════════════════════════════════════════════════════════════════════════════

heading1("5.4 FURTHER WORKS")

body(
    "The first suggestion for further work is to replace the SHAP KernelExplainer with a "
    "gradient-based attribution method such as Integrated Gradients or DeepLIFT, which are "
    "natively compatible with PyTorch transformer models and do not require repeated forward "
    "passes for each explanation. This would reduce the SHAP computation time from "
    "approximately 30 seconds to under one second, making real-time explanations feasible. "
    "This improvement would directly address the computational cost limitation identified in "
    "the previous section and would significantly enhance the user experience of the system."
)

body(
    "The second suggestion is to improve the tabular-to-text conversion template by replacing "
    "it with a generative Large Language Model capable of producing contextually richer and "
    "region-aware field descriptions. Such a template could incorporate local soil "
    "classifications, seasonal agricultural calendars, and crop rotation history, resulting "
    "in more informative input representations for the classifier. This would also improve "
    "the quality and naturalness of the attention-based reasoning outputs, as the model would "
    "attend to a more diverse vocabulary of agronomic concepts rather than a fixed set of "
    "descriptors."
)

body(
    "The third suggestion is to integrate live data acquisition from IoT soil sensors and "
    "public weather APIs such as OpenWeatherMap or NASA POWER. This would remove the "
    "dependency on manual parameter entry and enable the system to automatically retrieve "
    "up-to-date field data and generate recommendations on a scheduled basis. Automated PSI "
    "drift alerts could be triggered when the live input data deviates significantly from "
    "the training distribution, transforming the system from a query-on-demand tool into a "
    "proactive field advisory platform suitable for continuous agricultural monitoring."
)

body(
    "The fourth suggestion is to expand the training dataset by incorporating region-specific "
    "agricultural data from local agricultural departments or international bodies such as "
    "the Food and Agriculture Organisation (FAO). The current dataset contains only 2,200 "
    "samples across 22 crop classes and does not capture regional soil variation or seasonal "
    "cycles. A larger and more diverse dataset would improve the model's ability to generalise "
    "to real-world agricultural conditions and would allow the system to be extended to "
    "support a broader range of crop types and geographic regions."
)

# ════════════════════════════════════════════════════════════════════════════════
# 5.5 SUMMARY
# ════════════════════════════════════════════════════════════════════════════════

heading1("5.5 SUMMARY")

body(
    "This chapter discussed this project's core strengths, limitations, and further works "
    "that can be made, while revisiting the project objectives that were set at the beginning "
    "of the project. The project successfully achieved all three objectives, including "
    "identifying and assessing suitable LLM and XAI techniques for the crop recommendation "
    "system, designing and developing the full system using BERT, DistilBERT, SHAP, and "
    "attention-based reasoning, and evaluating both the functionality and the performance of "
    "the system. The system's key strengths are its dual-layer explainability capability, "
    "which provides both numerical SHAP attributions and natural language reasoning to "
    "support transparent decision-making, and its comprehensive admin portal that enables "
    "full model lifecycle management without code changes."
)

body(
    "However, the project also has some limitations, such as the high computational cost of "
    "the SHAP KernelExplainer, the fixed tabular-to-text conversion template that does not "
    "account for regional agricultural norms, and the reliance on manual parameter entry due "
    "to the absence of real-time sensor integration. To address these limitations, suggestions "
    "for further works have been provided, including replacing the KernelExplainer with a "
    "gradient-based attribution method, improving the text conversion template using a "
    "generative LLM, integrating live IoT sensor and weather API data, and expanding the "
    "training dataset with region-specific agricultural records. The CropAI system has the "
    "potential to become a more practical and widely accessible tool for agricultural "
    "decision-making if these limitations are addressed, which can contribute to more "
    "informed and data-driven crop planning for farmers."
)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r"e:\SEM6\Local_FYP_SHAP\Chapter5_Conclusion_and_Future_Work.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
