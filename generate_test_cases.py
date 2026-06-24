"""
Generates Section 4.5 Functional Testing as a standalone .docx file.
Run: python generate_test_cases.py
Output: Chapter4_FunctionalTesting.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Style helpers ─────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=10, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, bold=True, size=14)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)


def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=True, size=13)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)


def heading3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=True, italic=True, size=12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)


def body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.5)


def table_caption(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=True, size=11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)


def shade_cell(cell, fill="D9D9D9"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_cell_vertical_align(cell, align="center"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def add_tc_table(caption, headers, rows, col_widths):
    """
    Build a test-case table.
    headers : list of column header strings
    rows    : list of row dicts with keys matching headers list indices
    col_widths : list of Inches values
    Center-only columns: index 0 (TC#) and last (Status).
    All others left-aligned.
    """
    table_caption(caption)
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    CENTER_COLS = {0, len(headers) - 1}

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(cell)
        set_cell_vertical_align(cell)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            # Steps column: split on newline and add each as a separate paragraph
            lines = str(val).split("\n")
            for li, line in enumerate(lines):
                if li == 0:
                    para = cell.paragraphs[0]
                else:
                    para = cell.add_paragraph()
                run = para.add_run(line)
                set_font(run, size=10)
                para.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if ci in CENTER_COLS
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                para.paragraph_format.space_after = Pt(0)
            set_cell_vertical_align(cell, "top")

    # Column widths
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            cell.width = Inches(col_widths[ci])

    doc.add_paragraph()


# ── Headers ───────────────────────────────────────────────────────────────────

HEADERS = ["Test ID", "Test Name", "Precondition", "Steps",
           "Expected Outcome", "Actual Result", "Status"]

# col widths must sum to ≤ 6.0" (page width minus 1.25" margins each side)
# 0.45 + 0.85 + 1.0 + 1.2 + 1.0 + 1.05 + 0.45 = 6.0
WIDTHS = [0.45, 0.85, 1.0, 1.2, 1.0, 1.05, 0.45]

# ════════════════════════════════════════════════════════════════════════════
# 4.5  Functional Testing
# ════════════════════════════════════════════════════════════════════════════

heading1("4.5 Functional Testing")

heading2("4.5.1 Testing Overview")

body(
    "Functional testing was conducted using black-box testing — each test case "
    "is evaluated based on inputs and observable outputs only, without examining "
    "internal code execution. The test environment consisted of the FastAPI backend "
    "running on port 8000, the React frontend on port 5173, and a local MongoDB "
    "instance. A total of 19 test cases were designed across four modules: "
    "Authentication, Prediction, XAI, and Admin. Both valid inputs (happy path) "
    "and invalid inputs (error handling) are covered."
)

# ── 4.5.2 Authentication ──────────────────────────────────────────────────────

heading2("4.5.2 Authentication Module")

add_tc_table(
    "Table 4.7: Authentication Test Cases (TC-01 to TC-06)",
    HEADERS,
    [
        [
            "TC-01",
            "Register with valid credentials",
            "System is running; user has no existing account",
            "1. Navigate to /register\n2. Enter new username, email, and password\n3. Click Register",
            "Account created; redirected to /login",
            "Account created in MongoDB; redirected to /login",
            "Pass",
        ],
        [
            "TC-02",
            "Register with duplicate email",
            "TC-01 completed; email already exists in database",
            "1. Navigate to /register\n2. Enter same email as TC-01\n3. Click Register",
            "Error: email already registered",
            "400 — \"Email already registered\" returned",
            "Pass",
        ],
        [
            "TC-03",
            "Login with valid credentials",
            "TC-01 completed; account exists",
            "1. Navigate to /login\n2. Enter correct email and password\n3. Click Login",
            "JWT issued; redirected to /dashboard",
            "JWT (24-hour expiry) issued; user landed on /dashboard",
            "Pass",
        ],
        [
            "TC-04",
            "Login with invalid password",
            "Account exists in database",
            "1. Navigate to /login\n2. Enter correct email, wrong password\n3. Click Login",
            "Error: invalid credentials",
            "401 — \"Invalid credentials\" returned",
            "Pass",
        ],
        [
            "TC-05",
            "Access protected route without token",
            "User is not logged in; no JWT in context",
            "1. Open browser\n2. Navigate directly to /dashboard",
            "Redirected to /login",
            "React route guard detected missing token; redirected to /login",
            "Pass",
        ],
        [
            "TC-06",
            "Logout",
            "User is logged in with valid JWT",
            "1. Click the logout button on Navbar",
            "Token cleared; redirected to home page",
            "JWT removed from context; navigated to /",
            "Pass",
        ],
    ],
    WIDTHS,
)

# ── 4.5.3 Prediction ──────────────────────────────────────────────────────────

heading2("4.5.3 Prediction Module")

add_tc_table(
    "Table 4.8: Prediction Test Cases (TC-07 to TC-09)",
    HEADERS,
    [
        [
            "TC-07",
            "Prediction with valid inputs",
            "User is logged in; active model is loaded",
            "1. Navigate to /dashboard\n2. Set N=90, P=40, K=40, pH=6.5, Temp=23, Humidity=85, Rainfall=215\n3. Click Get Recommendation",
            "Recommended crop with confidence score and top-5 list returned",
            "Rice returned at 94.72% confidence with full top-5 breakdown",
            "Pass",
        ],
        [
            "TC-08",
            "Prediction at minimum slider values",
            "User is logged in; active model is loaded",
            "1. Navigate to /dashboard\n2. Set all 7 sliders to minimum values\n3. Click Get Recommendation",
            "Valid crop returned without error",
            "Prediction returned successfully without error",
            "Pass",
        ],
        [
            "TC-09",
            "Prediction at maximum slider values",
            "User is logged in; active model is loaded",
            "1. Navigate to /dashboard\n2. Set all 7 sliders to maximum values\n3. Click Get Recommendation",
            "Valid crop returned without error",
            "Prediction returned successfully without error",
            "Pass",
        ],
    ],
    WIDTHS,
)

# ── 4.5.4 XAI ────────────────────────────────────────────────────────────────

heading2("4.5.4 XAI Module")

add_tc_table(
    "Table 4.9: XAI Test Cases (TC-10 to TC-11)",
    HEADERS,
    [
        [
            "TC-10",
            "SHAP explanation for valid input",
            "TC-07 completed; prediction result displayed on Dashboard",
            "1. Use TC-07 inputs on Dashboard\n2. Click Explain button\n3. Wait for SHAP computation (~30s)",
            "SHAP bar chart rendered with 7 feature contribution values and base value",
            "SHAP values returned for all 7 features; chart rendered correctly",
            "Pass",
        ],
        [
            "TC-11",
            "Attention-based reasoning for valid input",
            "TC-07 completed; prediction result displayed on Dashboard",
            "1. Use TC-07 inputs on Dashboard\n2. Click Reasoning button",
            "Natural language sentence generated with top-3 features identified",
            "Reasoning sentence generated and displayed on Dashboard",
            "Pass",
        ],
    ],
    WIDTHS,
)

# ── 4.5.5 Admin ───────────────────────────────────────────────────────────────

heading2("4.5.5 Admin Module")

add_tc_table(
    "Table 4.10: Admin Test Cases (TC-12 to TC-19)",
    HEADERS,
    [
        [
            "TC-12",
            "Admin login with admin credentials",
            "Server started; default admin seeded at startup",
            "1. Navigate to /admin/login\n2. Enter admin@cropai.com and admin123\n3. Click Login",
            "Admin JWT (role=admin) issued; redirected to /admin/users",
            "JWT with admin role issued; redirected to Admin Users Page",
            "Pass",
        ],
        [
            "TC-13",
            "Admin route with regular user token",
            "Regular user is logged in with valid JWT",
            "1. Use regular user JWT to call an admin API endpoint",
            "403 Forbidden returned",
            "403 — \"Admin access required\" returned",
            "Pass",
        ],
        [
            "TC-14",
            "Toggle user active status",
            "Admin is logged in; TC-01 user account exists",
            "1. Navigate to /admin/users\n2. Click Disable on TC-01 user\n3. Attempt login as TC-01",
            "User disabled; login attempt rejected",
            "Status toggled in database; TC-01 login returned 401",
            "Pass",
        ],
        [
            "TC-15",
            "Delete user account",
            "Admin is logged in; TC-01 user account exists",
            "1. Navigate to /admin/users\n2. Click Delete on TC-01 user\n3. Confirm deletion",
            "User and all prediction history removed from database",
            "User deleted from MongoDB; prediction history cleared",
            "Pass",
        ],
        [
            "TC-16",
            "Switch active model",
            "Admin is logged in; both BERT and DistilBERT are loaded",
            "1. Navigate to /admin/models\n2. Click Set Active on DistilBERT\n3. Make a new prediction",
            "DistilBERT used for subsequent predictions",
            "active_model.json updated; next prediction returned model_used = DISTILBERT",
            "Pass",
        ],
        [
            "TC-17",
            "Trigger model retrain",
            "Admin is logged in; training dataset is available",
            "1. Navigate to /admin/models\n2. Click Retrain BERT",
            "Background retrain starts; progress updates displayed in UI",
            "Retrain thread started; stages shown (Loading → Training → Evaluating)",
            "Pass",
        ],
        [
            "TC-18",
            "Drift report with sufficient data",
            "Admin is logged in; ≥ 20 predictions exist in database",
            "1. Navigate to /admin/drift",
            "PSI report with all 7 features and daily confidence trend displayed",
            "PSI values computed and rendered with correct status indicators",
            "Pass",
        ],
        [
            "TC-19",
            "Drift report with insufficient data",
            "Admin is logged in; fewer than 20 predictions in database",
            "1. Navigate to /admin/drift",
            "Warning: insufficient data for PSI analysis",
            "Backend returned min_samples flag; UI displayed insufficient data notice",
            "Pass",
        ],
    ],
    WIDTHS,
)

# ── 4.5.6 Summary ─────────────────────────────────────────────────────────────

heading2("4.5.6 Testing Summary")

body(
    "Table 4.11 summarises the test results across all four modules. All 19 test "
    "cases passed, confirming that the system correctly handles both valid inputs "
    "and invalid inputs across every tested module. The 100% pass rate confirms "
    "that Research Objective 3 is met — the system functions correctly across all "
    "tested scenarios."
)

table_caption("Table 4.11: Test Results Summary by Module")
summary_t = doc.add_table(rows=6, cols=5)
summary_t.style = "Table Grid"
summary_t.alignment = WD_TABLE_ALIGNMENT.CENTER

summary_headers = ["Module", "Test Cases", "Passed", "Failed", "Pass Rate"]
summary_rows = [
    ["Authentication", "TC-01 to TC-06", "6", "0", "100%"],
    ["Prediction",     "TC-07 to TC-09", "3", "0", "100%"],
    ["XAI",           "TC-10 to TC-11", "2", "0", "100%"],
    ["Admin",         "TC-12 to TC-19", "8", "0", "100%"],
    ["Total",         "19",             "19","0", "100%"],
]
summary_widths = [1.4, 1.4, 0.9, 0.9, 0.9]

hrow = summary_t.rows[0]
for i, h in enumerate(summary_headers):
    cell = hrow.cells[i]
    cell.text = ""
    run = cell.paragraphs[0].add_run(h)
    set_font(run, bold=True, size=10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    shade_cell(cell)

for ri, row_data in enumerate(summary_rows):
    row = summary_t.rows[ri + 1]
    is_total = ri == 4
    for ci, val in enumerate(row_data):
        cell = row.cells[ci]
        cell.text = ""
        run = cell.paragraphs[0].add_run(val)
        set_font(run, bold=is_total, size=10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if is_total:
            shade_cell(cell, "F2F2F2")

for row in summary_t.rows:
    for ci, cell in enumerate(row.cells):
        cell.width = Inches(summary_widths[ci])

doc.add_paragraph()

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r"e:\SEM6\Local_FYP_SHAP\Chapter4_FunctionalTesting.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
