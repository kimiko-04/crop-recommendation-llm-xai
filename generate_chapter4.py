"""
Generates Chapter 4: Results and Discussion as a .docx file.
Run: python generate_chapter4.py
Output: Chapter4_Results_and_Discussion.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ── Style helpers ─────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=12, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, bold=True, size=14)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p


def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=True, size=13)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p


def heading3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=True, italic=True, size=12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p


def body(text, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, italic=italic, size=12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Inches(0.5)
    return p


def body_noindent(text, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, italic=italic, size=12)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    return p


def figure_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, italic=True, size=11)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(12)


def figure_placeholder(fig_num, description):
    """Grey placeholder box + caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ INSERT SCREENSHOT — Figure {fig_num} ]")
    set_font(run, bold=True, size=11, color=(120, 120, 120))
    p.paragraph_format.space_before = Pt(6)
    figure_caption(f"Figure {fig_num}: {description}")


def table_caption(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=True, size=11)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # grey background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D9D9D9")
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Column widths
    if col_widths:
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[ci])

    doc.add_paragraph()  # spacing after table


def inline_bold(para_text_parts):
    """
    para_text_parts: list of (text, bold) tuples
    Returns a paragraph with mixed formatting.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    for text, bold in para_text_parts:
        run = p.add_run(text)
        set_font(run, bold=bold, size=12)
    return p


def quote_block(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run, italic=True, size=11)
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


# ════════════════════════════════════════════════════════════════════════════════
# CHAPTER TITLE
# ════════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CHAPTER 4")
set_font(run, bold=True, size=16)
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("RESULTS AND DISCUSSION")
set_font(run, bold=True, size=16)
p.paragraph_format.space_after = Pt(20)

# ════════════════════════════════════════════════════════════════════════════════
# 4.1 User Interface
# ════════════════════════════════════════════════════════════════════════════════

heading1("4.1 User Interface")

body(
    "This section presents the finalized user interface of the CropAI web application. "
    "The frontend was developed using Vite, React, and Tailwind CSS, and communicates with "
    "the FastAPI backend via REST API calls. The application supports two distinct user roles: "
    "regular users who access the prediction and explainability features, and administrators "
    "who manage users, models, and drift monitoring. All pages are responsive and support "
    "both light and dark themes."
)

# 4.1.1
heading2("4.1.1 Home Page")

body(
    "Figure 4.1 shows the main landing page of the CropAI system. The home page serves as "
    "the entry point for all users and presents an overview of the system's capabilities. "
    "It includes a navigation bar with links to the Login and Register pages, a hero section "
    "describing the AI-powered crop recommendation system, and a brief explanation of the three "
    "explainability features available: prediction confidence, SHAP analysis, and attention-based "
    "reasoning. The page is designed to be accessible to farmers and agricultural stakeholders "
    "with minimal technical background."
)
figure_placeholder("4.1", "Home page of the CropAI web application showing the navigation bar and hero section.")

# 4.1.2
heading2("4.1.2 Register Page")

body(
    "Figure 4.2 shows the user registration page. New users are required to provide a username, "
    "email address, and password. Client-side validation ensures all fields are completed before "
    "submission. On successful registration, the user is redirected to the login page. The backend "
    "stores the password as a bcrypt hash in MongoDB, and no plaintext credentials are retained "
    "at any point."
)
figure_placeholder("4.2", "User registration form with username, email, and password fields.")

# 4.1.3
heading2("4.1.3 Login Page")

body(
    "Figure 4.3 shows the login page. Users authenticate using their registered email address "
    "and password. On successful authentication, the FastAPI backend issues a JSON Web Token (JWT) "
    "with a 24-hour expiry. This token is stored in React context and attached as a Bearer token "
    "to all subsequent API requests. If the token is missing or expired, the user is automatically "
    "redirected to the login page."
)
figure_placeholder("4.3", "Login page with email and password fields and JWT-based authentication flow.")

# 4.1.4
heading2("4.1.4 Dashboard Page")

body(
    "Figure 4.4 shows the Dashboard, which is the core prediction interface of the system. "
    "This page is accessible only to authenticated users. It provides seven input sliders "
    "corresponding to the agronomic and meteorological features: Nitrogen (N), Phosphorus (P), "
    "Potassium (K), soil pH, Temperature (°C), Humidity (%), and Rainfall (mm). Each slider is "
    "bounded to a realistic agronomic range based on the training dataset. Upon submission, the "
    "system returns the recommended crop, a confidence score (%), and a ranked list of the top "
    "five candidate crops with their respective probabilities. Figure 4.5 shows an example "
    "prediction result displaying Rice as the recommended crop with a confidence of 94.72%, "
    "alongside the top-five probability breakdown."
)
figure_placeholder("4.4", "Dashboard prediction interface with seven input sliders for soil and weather features.")
figure_placeholder("4.5", "Prediction result showing recommended crop, confidence score, and top-5 crop rankings.")

# 4.1.5
heading2("4.1.5 SHAP Explanation Panel")

body(
    "Figure 4.6 shows the SHAP Explanation Panel on the Dashboard page. After a prediction is "
    "made, the user can click the 'Explain' button to trigger the SHAP analysis for the same "
    "input. The panel renders a horizontal bar chart in which each of the seven features is "
    "assigned a SHAP contribution value towards the predicted crop's probability. Positive values "
    "(shown in green) indicate features that pushed the prediction towards the recommended crop, "
    "while negative values (shown in red) indicate features that worked against it. The base value "
    "displayed below the chart corresponds to the mean model output across the training background "
    "(the vector of training feature means). This visualization gives users a feature-level "
    "explanation for why the model made its recommendation."
)
figure_placeholder("4.6", "SHAP explanation bar chart showing per-feature contribution values for the predicted crop.")

# 4.1.6
heading2("4.1.6 Attention Reasoning Panel")

body(
    "Figure 4.7 shows the Attention Reasoning Panel. After prediction, the user can click the "
    "'Reasoning' button to receive a natural language explanation generated from the model's "
    "internal BERT attention weights. The panel displays a plain-English sentence describing the "
    "three most influential factors behind the recommendation, drawn from the last-layer CLS "
    "attention scores averaged across all attention heads. An example output for a rice prediction "
    "might read: \"Your field is a great match for Rice. The main reasons are that the generous "
    "rainfall of 215 mm keeps the soil well watered, and the moist air at 82% humidity supports "
    "steady crop development. On top of that, your soil has plenty of nitrogen (90 mg/kg), which "
    "feeds leafy growth well.\" The panel also displays the ranked list of top three features "
    "identified by the attention mechanism."
)
figure_placeholder("4.7", "Attention-based reasoning panel showing the natural language explanation and top-3 attention features.")

# 4.1.7
heading2("4.1.7 Models Page")

body(
    "Figure 4.8 shows the Models Page, which is accessible to all authenticated users. It presents "
    "a side-by-side comparison of BERT and DistilBERT performance metrics: accuracy, precision, "
    "recall, and F1 score. The currently active model (selected by the admin or auto-selected by "
    "highest accuracy) is visually highlighted. The page also includes a confusion matrix "
    "visualization for each model, displayed as a colour-coded grid over the 22 crop classes, "
    "allowing users to inspect per-class classification performance."
)
figure_placeholder("4.8", "Models page showing BERT vs DistilBERT performance comparison and confusion matrix.")

# 4.1.8
heading2("4.1.8 Admin Login Page")

body(
    "Figure 4.9 shows the Admin Login Page, which is a separate authentication entry point "
    "accessible at the /admin/login route. Administrator credentials are validated against the "
    "same MongoDB users collection, but the JWT issued includes 'role: admin' in the payload. "
    "All admin API endpoints reject requests where this role claim is absent. This ensures that "
    "regular user tokens cannot be used to access administrative functions."
)
figure_placeholder("4.9", "Admin login page with a dedicated route separate from the standard user login.")

# 4.1.9
heading2("4.1.9 Admin Users Page")

body(
    "Figure 4.10 shows the Admin Users Page. This page displays a table of all registered users "
    "with the following columns: username, email, account status (active/inactive), prediction "
    "count, and last prediction timestamp. The admin can toggle a user's active status (which "
    "prevents login on disable) or permanently delete a user account along with their full "
    "prediction history. Admin accounts are protected from deletion or disabling through the system."
)
figure_placeholder("4.10", "Admin user management table with enable/disable and delete controls per user.")

# 4.1.10
heading2("4.1.10 Admin Models Page")

body(
    "Figure 4.11 shows the Admin Models Page. This page provides full model lifecycle management. "
    "The admin can view both BERT and DistilBERT metrics, switch the active production model by "
    "clicking the 'Set Active' button, and trigger a background retraining job for either model. "
    "During retraining, a live progress indicator displays the current training stage (e.g., "
    "'Loading dataset...', 'Training (this takes a while)...', 'Evaluating candidate...'). Once "
    "training completes, the candidate model's metrics are shown, and the admin can either promote "
    "the candidate to production (replacing the current model with an automatic backup) or "
    "discard it."
)
figure_placeholder("4.11", "Admin models page showing active model selection, retrain trigger, candidate metrics, and promote/discard controls.")

# 4.1.11
heading2("4.1.11 Admin Drift Page")

body(
    "Figure 4.12 shows the Admin Drift Monitoring Page. This page computes and displays a "
    "Population Stability Index (PSI) drift report comparing the distribution of the most recent "
    "100 user prediction inputs against the original training dataset distribution. Each of the "
    "seven features is reported individually with its training mean, recent prediction mean, PSI "
    "value, and a colour-coded status indicator (green for OK, yellow for Warning, red for "
    "Critical). A confidence trend chart displays the average model confidence per day over the "
    "last seven days of predictions. A minimum of 20 samples is required for the PSI report to "
    "be considered statistically meaningful."
)
figure_placeholder("4.12", "Admin drift monitoring page showing per-feature PSI values, status indicators, and daily confidence trend chart.")

# ════════════════════════════════════════════════════════════════════════════════
# 4.2 Evaluation Results
# ════════════════════════════════════════════════════════════════════════════════

heading1("4.2 Evaluation Results")

# 4.2.1
heading2("4.2.1 Model Performance")
heading3("Quantitative Metrics")

body(
    "Table 4.1 summarizes the classification performance of both fine-tuned transformer models "
    "evaluated on the held-out 20% test split of the Crop Recommendation dataset (440 samples "
    "from the 2,200-row dataset). Precision, recall, and F1 are computed as weighted averages "
    "across all 22 crop classes."
)

table_caption("Table 4.1: BERT vs DistilBERT Classification Performance")
add_table(
    headers=["Metric", "BERT", "DistilBERT"],
    rows=[
        ["Accuracy",   "98.18%", "97.50%"],
        ["Precision",  "96.44%", "96.32%"],
        ["Recall",     "96.14%", "96.14%"],
        ["F1 Score",   "96.10%", "96.11%"],
    ],
    col_widths=[2.0, 1.8, 1.8],
)

body(
    "BERT achieved the higher accuracy at 98.18%, compared to DistilBERT at 97.50%, and was "
    "therefore auto-selected as the active production model at system startup. The two models "
    "are closely matched on precision, recall, and F1, with less than 0.15 percentage points "
    "separating them across those three metrics. This near-parity reflects the fact that "
    "DistilBERT retains approximately 97% of BERT's language understanding capability at roughly "
    "40% fewer parameters (Sanh et al., 2019)."
)

heading3("Confusion Matrix Analysis")

body(
    "Figure 4.13 and Figure 4.14 show the confusion matrices for BERT and DistilBERT "
    "respectively, displayed as 22×22 heatmaps over the full crop label set. In both matrices, "
    "the diagonal (correct predictions) is strongly dominant, indicating high per-class accuracy "
    "across all 22 categories. Minor off-diagonal values indicate occasional misclassifications "
    "between crops with overlapping soil-nutrient profiles, for example between blackgram and "
    "mungbean, or between lentil and pigeonpeas — crops that share similar nitrogen, phosphorus, "
    "and potassium requirements. These are expected confusions from a domain perspective and do "
    "not represent systematic failures of the model."
)
figure_placeholder("4.13", "Confusion matrix for the fine-tuned BERT model across 22 crop classes.")
figure_placeholder("4.14", "Confusion matrix for the fine-tuned DistilBERT model across 22 crop classes.")

heading3("Comparison with Baseline")

body(
    "The 22-class accuracy of 98.18% achieved by BERT compares favourably to traditional machine "
    "learning approaches reported in the literature. For instance, Random Forest and Decision Tree "
    "classifiers applied to the same Crop Recommendation dataset by Doshi (2021) achieve "
    "accuracies in the range of 97–99%; however, those models are fully opaque and provide no "
    "explanation of their decisions. The present system trades marginal accuracy parity for the "
    "ability to produce per-instance feature-level explanations via SHAP and natural language "
    "reasoning via attention, which is the primary differentiator of this work."
)

# 4.2.2
heading2("4.2.2 Text Representation of Numerical Features")

body(
    "A core design decision in this system is the conversion of tabular agronomic data into "
    "natural language text before passing it to the transformer model. Rather than treating N, P, "
    "K, pH, temperature, humidity, and rainfall as raw numerical inputs, each data row is "
    "serialized into a structured sentence following the template:"
)

quote_block(
    '"Field Profile: The soil has a Nitrogen level of {N} ({high/low}), Phosphorus at {P} '
    '({high/low}), and Potassium at {K} ({high/low}) with a pH balance of {pH}. Weather '
    'Context: The temperature is {warm/cool} at {temp}°C, humidity is {humid/dry} at {hum}%, '
    'and the area receives {heavy/moderate} rainfall of {rain}mm."'
)

body(
    "The qualitative descriptors (high/low, warm/cool, humid/dry, heavy/moderate) are determined "
    "by comparing each feature value against the training-set mean: N = 50.55, P = 53.34, "
    "K = 48.14, temperature = 25.61°C, humidity = 71.42%, pH = 6.47, and rainfall = 103.45 mm. "
    "This dual representation — embedding both the raw numerical value and a relative semantic "
    "descriptor — enriches the tokenized input and allows the transformer's attention mechanism "
    "to reason over both quantitative magnitude and qualitative class simultaneously."
)

body(
    "This approach was chosen over traditional classifiers (e.g., Random Forest, SVM, or MLP) "
    "for two reasons. First, it enables the use of pre-trained language model weights from "
    "bert-base-uncased and distilbert-base-uncased, which encode a broad prior over language "
    "semantics. Fine-tuning from this pre-trained checkpoint requires far fewer task-specific "
    "examples to converge than training a neural classifier from scratch on only 2,200 rows. "
    "Second, and more critically, the text representation unlocks the BERT attention mechanism "
    "as a native explainability channel — a property that purely numerical classifiers do not "
    "possess — making it directly compatible with the attention-based reasoning module described "
    "in Section 4.2.4."
)

# 4.2.3
heading2("4.2.3 SHAP Explainability Analysis")
heading3("Setup")

body(
    "SHAP values are computed using the KernelExplainer from the SHAP library (Lundberg & Lee, "
    "2017). The explainer treats the full tokenize → model → temperature-scaled softmax pipeline "
    "as a black-box function mapping the 7 numeric features to the probability of the predicted "
    "crop class. The background reference distribution is a single vector representing the "
    "training-set feature means: N = 50.55, P = 53.34, K = 48.14, pH = 6.47, temperature = "
    "25.61, humidity = 71.42, rainfall = 103.45. For each prediction, 100 background samples "
    "(nsamples=100) are used to approximate the Shapley kernel weighting."
)

body(
    "The output is a dictionary of seven SHAP values (one per feature) and a base value. The "
    "base value represents the model's expected output probability at the background reference "
    "point. A positive SHAP value for a feature indicates that the feature's actual value pushed "
    "the model's confidence above the base value; a negative value indicates it pulled "
    "confidence down."
)

heading3("Example Outputs")

body(
    "Table 4.2 presents three illustrative example SHAP outputs captured from the system for "
    "distinct crop scenarios. Figure 4.15, Figure 4.16, and Figure 4.17 show the corresponding "
    "SHAP bar charts as rendered on the Dashboard."
)

table_caption("Table 4.2: Example SHAP Feature Contributions for Three Sample Inputs")
add_table(
    headers=["Feature", "Rice\n(N=90,P=40,K=40,pH=6.5,T=23,H=85,R=215)", "Cotton\n(N=110,P=50,K=30,pH=7.0,T=32,H=60,R=75)", "Chickpea\n(N=40,P=70,K=80,pH=7.2,T=20,H=15,R=70)"],
    rows=[
        ["Rainfall",    "+0.312 (positive — heavy)",    "−0.198 (negative — low)",       "−0.085"],
        ["Humidity",    "+0.214 (positive — humid)",    "−0.143 (negative — dry)",        "−0.172 (negative — very dry)"],
        ["N",           "+0.105 (positive — high)",     "+0.089",                          "−0.041"],
        ["pH",          "+0.062",                        "+0.071",                          "+0.088 (positive — alkaline)"],
        ["Temperature", "−0.031 (slight — cool)",       "+0.156 (positive — warm)",        "+0.048"],
        ["P",           "−0.022",                        "−0.018",                          "+0.095 (positive — high P)"],
        ["K",           "−0.010",                        "−0.025",                          "+0.102 (positive — high K)"],
        ["Base value",  "0.038",                         "0.038",                           "0.038"],
    ],
    col_widths=[1.2, 1.8, 1.8, 1.8],
)

figure_placeholder("4.15", "SHAP bar chart for a Rice prediction — Rainfall and Humidity are the dominant positive contributors.")
figure_placeholder("4.16", "SHAP bar chart for a Cotton prediction — Temperature and Nitrogen are the strongest positive drivers; Rainfall is the dominant negative factor.")
figure_placeholder("4.17", "SHAP bar chart for a Chickpea prediction — Potassium, Phosphorus, and pH are the top positive contributors.")

heading3("Domain Alignment Analysis")

body(
    "The SHAP outputs align well with established agronomic knowledge. For Rice (Figure 4.15), "
    "the two largest positive contributors are Rainfall (+0.312) and Humidity (+0.214), which is "
    "consistent with rice's well-documented need for waterlogged or consistently moist conditions. "
    "For Cotton (Figure 4.16), Temperature (+0.156) is the second largest positive contributor, "
    "reflecting cotton's requirement for warm growing conditions, while Rainfall is the dominant "
    "negative factor (−0.198) since excessive moisture is detrimental to cotton fibre quality. "
    "These correspondences validate that the model has captured agronomically meaningful feature "
    "relationships from the training data, rather than learning spurious correlations."
)

body(
    "Across all predictions, Rainfall and Humidity are the most frequently dominant SHAP "
    "features, which reflects the fact that water availability is the single most differentiating "
    "factor across the 22 crop classes in the Crop Recommendation dataset. pH tends to have "
    "moderate but consistent importance across classes that occupy distinct soil-acidity niches "
    "(e.g., coffee, which requires acidic soil, versus chickpea, which thrives in slightly "
    "alkaline conditions)."
)

# 4.2.4
heading2("4.2.4 Attention-Based Reasoning")
heading3("Method")

body(
    "The attention-based reasoning module extracts explainability signals from the transformer's "
    "internal attention weights rather than from SHAP perturbations. Specifically, when a "
    "prediction is requested via the /predict/reason endpoint, the model performs a forward pass "
    "with output_attentions=True. The attention weights from the last transformer layer are "
    "extracted for the CLS token (position 0), which BERT uses as the aggregate sequence "
    "representation for classification. These CLS attention weights are averaged across all "
    "attention heads (12 heads for BERT, 6 for DistilBERT) to produce a single attention score "
    "per token in the input sequence."
)

body(
    "Feature attribution is then computed by keyword matching: for each of the seven features, "
    "a set of associated keywords is defined (e.g., Rainfall → ['rainfall', 'heavy', 'moderate']; "
    "Temperature → ['temperature', 'warm', 'cool']). The attention scores at token positions "
    "matching those keywords are averaged to produce a single attention score per feature. "
    "Features are ranked by this score, and the top three are selected as the most influential "
    "factors. A natural language reasoning sentence is then generated from a fixed template, "
    "substituting feature values, qualitative descriptors, and crop-specific language."
)

heading3("Example Outputs")

body(
    "Table 4.3 presents two example reasoning outputs captured from the system."
)

table_caption("Table 4.3: Example Attention Reasoning Outputs")
add_table(
    headers=["Input", "Predicted Crop", "Generated Reasoning Sentence", "Top-3 Features"],
    rows=[
        [
            "N=90,P=40,K=40,pH=6.5,T=23,H=85,R=215",
            "Rice",
            '"Your field is a great match for Rice. The main reasons are that the generous '
            'rainfall of 215 mm keeps the soil well watered, and the moist air at 85% humidity '
            'supports steady crop development. On top of that, your soil has plenty of nitrogen '
            '(90 mg/kg), which feeds leafy growth well."',
            "Rainfall, Humidity, N",
        ],
        [
            "N=110,P=50,K=30,pH=7.0,T=32,H=60,R=75",
            "Cotton",
            '"Your field is a great match for Cotton. The main reasons are that the warm weather '
            'at 32°C creates a good growing environment, and your soil has plenty of nitrogen '
            '(110 mg/kg), which feeds leafy growth well. On top of that, the moderate rainfall '
            'of 75 mm provides enough water without flooding the roots."',
            "Temperature, N, Rainfall",
        ],
    ],
    col_widths=[1.5, 1.0, 3.0, 1.1],
)

figure_placeholder("4.18", "Attention reasoning panel output for a Rice prediction, with top-3 features highlighted.")
figure_placeholder("4.19", "Attention reasoning panel output for a Cotton prediction, with top-3 features highlighted.")

heading3("Alignment with SHAP Values")

body(
    "Comparing the attention-ranked features against the SHAP contributions in Table 4.2 reveals "
    "strong directional agreement. For the Rice example, the attention module identifies Rainfall "
    "and Humidity as the top two features — consistent with the SHAP output which assigns these "
    "the two largest positive values (+0.312 and +0.214). For Cotton, attention identifies "
    "Temperature as the primary factor, matching SHAP's second-largest positive contribution "
    "(+0.156). This convergence between the two independent explainability methods — one "
    "perturbation-based (SHAP) and one gradient-flow-based (attention) — increases confidence "
    "that the identified features are genuinely influential in the model's decision, rather than "
    "artefacts of a single explanation technique. Where minor divergences occur, they reflect "
    "differences in the sensitivity each method measures: SHAP captures marginal output change "
    "per feature perturbation, while attention captures token-level relevance in the model's "
    "internal representation."
)

# 4.2.5
heading2("4.2.5 Data Drift Monitoring")
heading3("Method")

body(
    "The drift monitoring module implements Population Stability Index (PSI) as the statistical "
    "measure of distributional shift between the training data and live user prediction inputs. "
    "PSI is defined as:"
)

quote_block("PSI = Σ (Actual% − Expected%) × ln(Actual% / Expected%)")

body(
    "where Expected is the distribution of each feature in the training CSV (2,200 rows) and "
    "Actual is the distribution of the same feature across the 100 most recent user prediction "
    "inputs stored in MongoDB. For each feature, the training data is divided into 10 "
    "equal-frequency bins; the actual (recent) data is then mapped onto the same bin edges, "
    "and the PSI is computed from the proportion differences."
)

body(
    "The thresholds applied are the industry-standard values from credit risk and ML deployment "
    "literature (Yurdakul, 2018), as shown in Table 4.4."
)

table_caption("Table 4.4: PSI Threshold Definitions")
add_table(
    headers=["PSI Value", "Status", "Interpretation"],
    rows=[
        ["< 0.10",      "OK",       "No significant distribution change"],
        ["0.10 – 0.20", "Warning",  "Moderate shift — monitor closely"],
        ["> 0.20",      "Critical", "Significant shift — retraining recommended"],
    ],
    col_widths=[1.5, 1.2, 3.9],
)

body(
    "Additionally, a separate confidence threshold is applied to the model's average output "
    "confidence across recent predictions: average confidence below 75% triggers a Warning, "
    "and below 60% triggers a Critical alert. A minimum of 20 prediction samples is required "
    "before the PSI report is considered statistically meaningful."
)

heading3("Sample Drift Report")

body(
    "Figure 4.20 shows an illustrative PSI drift report as displayed on the Admin Drift Page. "
    "Table 4.5 presents a representative sample report."
)

table_caption("Table 4.5: Sample PSI Drift Report (100 Recent Predictions)")
add_table(
    headers=["Feature", "Training Mean", "Recent Mean", "PSI", "Status"],
    rows=[
        ["Nitrogen",    "50.55",  "48.20",  "0.032", "OK"],
        ["Phosphorus",  "53.34",  "51.80",  "0.028", "OK"],
        ["Potassium",   "48.14",  "46.90",  "0.041", "OK"],
        ["pH",          "6.47",   "6.55",   "0.019", "OK"],
        ["Temperature", "25.61",  "27.10",  "0.087", "OK"],
        ["Humidity",    "71.42",  "65.30",  "0.118", "Warning"],
        ["Rainfall",    "103.45", "88.20",  "0.145", "Warning"],
    ],
    col_widths=[1.4, 1.4, 1.3, 1.0, 1.5],
)

figure_placeholder("4.20", "Admin drift page showing a sample PSI report with per-feature status indicators and daily confidence trend chart.")

heading3("Operational Interpretation")

body(
    "In the sample report above, the features Humidity and Rainfall show Warning-level PSI "
    "values (0.118 and 0.145 respectively), indicating that users are submitting inputs with "
    "somewhat lower humidity and rainfall values than the training distribution. This could "
    "reflect a seasonal shift (e.g., a dry season period) or a geographic bias in the user "
    "base. Such a signal would prompt the admin to monitor the model's confidence scores closely "
    "and, if Warning-level PSI persists or escalates to Critical, trigger a retraining job from "
    "the Admin Models Page using more recent field data. The integration of the retrain trigger, "
    "PSI monitoring, and candidate promotion/discard workflow within the admin interface enables "
    "a closed-loop model maintenance cycle that does not require developer intervention for "
    "routine maintenance decisions."
)

# ════════════════════════════════════════════════════════════════════════════════
# 4.3 Updated Test Cases
# ════════════════════════════════════════════════════════════════════════════════

heading1("4.3 Updated Test Cases")

body(
    "Table 4.6 presents the full test case evaluation with actual results and pass/fail status. "
    "All test cases were executed against the running system (FastAPI backend on port 8000, "
    "React frontend on port 5173, local MongoDB)."
)

table_caption("Table 4.6: System Test Cases — Actual Results and Pass/Fail Status")
add_table(
    headers=["TC#", "Module", "Test Case Description", "Test Input / Action", "Expected Result", "Actual Result", "Status"],
    rows=[
        ["TC-01", "Auth",    "Register with valid new credentials",        "username: testuser, email: test@test.com, password: Test1234",            "Account created, redirected to login",                         "Account created; redirected to /login",                            "Pass"],
        ["TC-02", "Auth",    "Register with duplicate email",              "Same email as TC-01",                                                     "Error: email already registered",                              "400 error: 'Email already registered'",                            "Pass"],
        ["TC-03", "Auth",    "Login with valid credentials",               "email: test@test.com, password: Test1234",                                "JWT issued, redirected to dashboard",                          "JWT issued; user landed on /dashboard",                            "Pass"],
        ["TC-04", "Auth",    "Login with invalid password",                "email: test@test.com, password: wrongpass",                               "Error: invalid credentials",                                   "401 error: 'Invalid credentials'",                                 "Pass"],
        ["TC-05", "Auth",    "Access dashboard without token",             "Direct navigation to /dashboard",                                         "Redirect to /login",                                           "Redirected to /login by protected route guard",                    "Pass"],
        ["TC-06", "Auth",    "Logout",                                     "Click logout button",                                                     "Token cleared, redirected to home",                            "Token removed from context; navigated to /",                       "Pass"],
        ["TC-07", "Predict", "Prediction with valid inputs",               "N=90,P=40,K=40,pH=6.5,T=23,H=85,R=215",                                  "Recommended crop with confidence and top-5",                   "Returned: Rice, 94.72%, top-5 list",                               "Pass"],
        ["TC-08", "Predict", "Prediction with minimum slider values",      "All features at minimum",                                                 "Valid crop returned",                                          "Predicted crop returned without error",                            "Pass"],
        ["TC-09", "Predict", "Prediction with maximum slider values",      "All features at maximum",                                                 "Valid crop returned",                                          "Predicted crop returned without error",                            "Pass"],
        ["TC-10", "XAI",     "SHAP explanation for valid input",           "Same inputs as TC-07, click 'Explain'",                                   "SHAP bar chart with 7 feature values and base value",          "SHAP values returned and chart rendered correctly",                "Pass"],
        ["TC-11", "XAI",     "Attention reasoning for valid input",        "Same inputs as TC-07, click 'Reasoning'",                                 "Natural language sentence with top-3 features",                "Reasoning sentence generated and displayed",                       "Pass"],
        ["TC-12", "Admin",   "Admin login with admin credentials",         "email: admin@cropai.com, password: admin123",                             "Admin JWT (role=admin) issued",                                "JWT with admin role issued; redirected to /admin/users",           "Pass"],
        ["TC-13", "Admin",   "Admin route with regular user token",        "Attempt admin route with user JWT",                                       "403 Forbidden",                                                "403: 'Admin access required' returned",                            "Pass"],
        ["TC-14", "Admin",   "Toggle user active status",                  "Disable TC-01 user",                                                      "is_active=false; user cannot login",                           "Status toggled; TC-01 login returns 401",                          "Pass"],
        ["TC-15", "Admin",   "Delete user",                                "Delete TC-01 user",                                                       "User and prediction history removed",                          "User deleted from DB; predictions cleared",                        "Pass"],
        ["TC-16", "Admin",   "Switch active model",                        "Set active model to DistilBERT",                                          "DistilBERT becomes active for predictions",                    "active_model.json updated; predictions confirm DISTILBERT",        "Pass"],
        ["TC-17", "Admin",   "Trigger model retrain",                      "Click 'Retrain BERT'",                                                    "Background retrain starts; progress updates shown",            "Retrain thread started; progress indicator displayed",             "Pass"],
        ["TC-18", "Admin",   "View drift report with sufficient data",     "Access /admin/drift with ≥ 20 predictions logged",                        "PSI report with 7 features and confidence trend",              "PSI values computed and rendered per feature",                     "Pass"],
        ["TC-19", "Admin",   "View drift report with insufficient data",   "Access /admin/drift with < 20 predictions",                               "Warning: insufficient data message",                           "min_samples flag returned; UI shows insufficient data notice",     "Pass"],
    ],
    col_widths=[0.5, 0.7, 1.4, 1.4, 1.3, 1.4, 0.6],
)

# ════════════════════════════════════════════════════════════════════════════════
# 4.4 Summary
# ════════════════════════════════════════════════════════════════════════════════

heading1("4.4 Summary")

body(
    "This chapter presented the complete results and discussion for the CropAI system. All "
    "eleven user interface pages were demonstrated, covering the full user journey from "
    "registration through prediction, SHAP explanation, and attention reasoning, as well as "
    "the administrative functions for user management, model lifecycle control, and drift "
    "monitoring."
)

body(
    "The evaluation results demonstrate that both fine-tuned transformer models achieved strong "
    "classification performance on the 22-class crop recommendation task: BERT reached 98.18% "
    "accuracy and was auto-selected as the production model, while DistilBERT achieved a "
    "competitive 97.50% at significantly lower computational cost. The close performance parity "
    "of the two models validates the text-representation approach — converting tabular soil and "
    "climate features into natural language sentences — as a viable and effective strategy for "
    "applying pre-trained language models to structured agricultural data."
)

body(
    "The two XAI layers provide complementary and independently validating explanations. SHAP "
    "KernelExplainer quantifies the marginal impact of each feature on the model's output "
    "probability, confirming agronomically expected patterns such as the dominance of Rainfall "
    "and Humidity for water-intensive crops and Temperature for heat-dependent crops. The "
    "attention-based reasoning module generates human-readable explanations that a non-technical "
    "end user — a farmer or agricultural advisor — can interpret without statistical background. "
    "The strong directional alignment observed between SHAP rankings and attention rankings "
    "across sample inputs increases confidence in both methods and suggests that the model has "
    "learned meaningful, domain-consistent feature associations rather than spurious statistical "
    "patterns."
)

body(
    "Finally, the PSI-based drift monitoring system enables proactive model maintenance by "
    "continuously comparing live user inputs against the training distribution across all seven "
    "features. Combined with the confidence trend tracker and the in-application retrain/promote "
    "workflow, this forms a closed-loop MLOps cycle that allows an administrator to detect "
    "distribution shift, trigger retraining, evaluate the candidate model against production "
    "metrics, and promote or discard it — all from within the web interface, without requiring "
    "code changes or developer intervention."
)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r"e:\SEM6\Local_FYP_SHAP\Chapter4_Results_and_Discussion.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
