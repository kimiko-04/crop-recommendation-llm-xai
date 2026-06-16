"""
Generates Chapter 4 Table of Contents, List of Figures, and List of Tables as a .docx file.
Run: python generate_toc.py
Output: Chapter4_TableOfContents.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)


def set_font(run, bold=False, italic=False, size=12, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    if color:
        run.font.color.rgb = RGBColor(*color)


def dot_leader_para(left_text, page, bold=False, italic=False, size=12, indent=0.0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(indent)

    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"),    "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"),    "8640")
    tabs.append(tab)
    pPr.append(tabs)

    run = p.add_run(left_text)
    set_font(run, bold=bold, italic=italic, size=size)

    tr = p.add_run("\t" + str(page))
    set_font(tr, bold=bold, italic=italic, size=size)
    return p


def section_title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=True, size=13)
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(8)


# ══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TABLE OF CONTENTS")
set_font(run, bold=True, size=14)
p.paragraph_format.space_after = Pt(16)

# Thin rule
p = doc.add_paragraph()
run = p.add_run("―" * 72)
set_font(run, size=10, color=(150, 150, 150))
p.paragraph_format.space_after = Pt(10)

# (level, text, page)
toc_entries = [
    (0, "CHAPTER 4: RESULTS AND DISCUSSION",                     1),
    (1, "4.1  User Interface",                                    1),
    (2, "4.1.1  Home Page",                                       1),
    (2, "4.1.2  Register Page",                                   2),
    (2, "4.1.3  Login Page",                                      2),
    (2, "4.1.4  Dashboard Page",                                  3),
    (2, "4.1.5  SHAP Explanation Panel",                          4),
    (2, "4.1.6  Attention Reasoning Panel",                       4),
    (2, "4.1.7  Models Page",                                     5),
    (2, "4.1.8  Admin Login Page",                                5),
    (2, "4.1.9  Admin Users Page",                                6),
    (2, "4.1.10  Admin Models Page",                              6),
    (2, "4.1.11  Admin Drift Page",                               7),
    (1, "4.2  Evaluation Results",                                7),
    (2, "4.2.1  Model Performance",                               7),
    (3, "Quantitative Metrics",                                    7),
    (3, "Confusion Matrix Analysis",                               8),
    (3, "Comparison with Baseline",                                9),
    (2, "4.2.2  Text Representation of Numerical Features",       9),
    (2, "4.2.3  SHAP Explainability Analysis",                   10),
    (3, "Setup",                                                  10),
    (3, "Example Outputs",                                        10),
    (3, "Domain Alignment Analysis",                              11),
    (2, "4.2.4  Attention-Based Reasoning",                      12),
    (3, "Method",                                                 12),
    (3, "Example Outputs",                                        12),
    (3, "Alignment with SHAP Values",                             13),
    (2, "4.2.5  Data Drift Monitoring",                          13),
    (3, "Method",                                                 13),
    (3, "Sample Drift Report",                                    14),
    (3, "Operational Interpretation",                             15),
    (1, "4.3  Updated Test Cases",                               15),
    (1, "4.4  Summary",                                          17),
]

indent_map = {0: 0.0, 1: 0.0, 2: 0.3, 3: 0.6}
size_map   = {0: 12,  1: 12,  2: 12,  3: 11}
bold_map   = {0: True, 1: True, 2: False, 3: False}
italic_map = {0: False, 1: False, 2: False, 3: True}

for level, text, page in toc_entries:
    dot_leader_para(
        left_text=text,
        page=page,
        bold=bold_map[level],
        italic=italic_map[level],
        size=size_map[level],
        indent=indent_map[level],
    )

# ══════════════════════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ══════════════════════════════════════════════════════════════════════════════

section_title("LIST OF FIGURES")

figures = [
    ("4.1",  "Home page of the CropAI web application showing navigation bar and hero section",            1),
    ("4.2",  "User registration form with username, email, and password fields",                            2),
    ("4.3",  "Login page with email and password fields and JWT-based authentication flow",                 2),
    ("4.4",  "Dashboard prediction interface with seven input sliders for soil and weather features",       3),
    ("4.5",  "Prediction result showing recommended crop, confidence score, and top-5 crop rankings",      3),
    ("4.6",  "SHAP explanation bar chart showing per-feature contribution values for the predicted crop",  4),
    ("4.7",  "Attention-based reasoning panel showing natural language explanation and top-3 features",     4),
    ("4.8",  "Models page showing BERT vs DistilBERT performance comparison and confusion matrix",          5),
    ("4.9",  "Admin login page with a dedicated route separate from the standard user login",              5),
    ("4.10", "Admin user management table with enable/disable and delete controls per user",               6),
    ("4.11", "Admin models page with active model selection, retrain, and promote/discard controls",       6),
    ("4.12", "Admin drift monitoring page with per-feature PSI values and confidence trend chart",         7),
    ("4.13", "Confusion matrix for the fine-tuned BERT model across 22 crop classes",                      8),
    ("4.14", "Confusion matrix for the fine-tuned DistilBERT model across 22 crop classes",               8),
    ("4.15", "SHAP bar chart for a Rice prediction — Rainfall and Humidity as dominant contributors",     10),
    ("4.16", "SHAP bar chart for a Cotton prediction — Temperature as strongest positive driver",         11),
    ("4.17", "SHAP bar chart for a Chickpea prediction — Potassium, Phosphorus, and pH as top features", 11),
    ("4.18", "Attention reasoning panel output for a Rice prediction with top-3 features highlighted",    12),
    ("4.19", "Attention reasoning panel output for a Cotton prediction with top-3 features highlighted",  12),
    ("4.20", "Admin drift page showing sample PSI report and daily confidence trend chart",               14),
]

for fig_num, desc, page in figures:
    dot_leader_para(
        left_text=f"Figure {fig_num}:   {desc}",
        page=page,
        bold=False,
        size=11,
        indent=0.0,
    )

# ══════════════════════════════════════════════════════════════════════════════
# LIST OF TABLES
# ══════════════════════════════════════════════════════════════════════════════

section_title("LIST OF TABLES")

tables_list = [
    ("4.1", "BERT vs DistilBERT Classification Performance",                     7),
    ("4.2", "Example SHAP Feature Contributions for Three Sample Inputs",        10),
    ("4.3", "Example Attention Reasoning Outputs",                               12),
    ("4.4", "PSI Threshold Definitions",                                         13),
    ("4.5", "Sample PSI Drift Report (100 Recent Predictions)",                  14),
    ("4.6", "System Test Cases — Actual Results and Pass/Fail Status",           15),
]

for tbl_num, desc, page in tables_list:
    dot_leader_para(
        left_text=f"Table {tbl_num}:   {desc}",
        page=page,
        bold=False,
        size=11,
        indent=0.0,
    )

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r"e:\SEM6\Local_FYP_SHAP\Chapter4_TableOfContents.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
